import os
from pypdf import PdfReader
from docx import Document as DocxDocument

def parse_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def parse_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text.strip()

def parse_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    text = []
    for para in doc.paragraphs:
        if para.text:
            text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                text.append(" | ".join(row_text))
    return "\n".join(text).strip()

def parse_document(file_path: str) -> str:
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".txt":
        return parse_txt(file_path)
    elif ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    else:
        # Fallback to plain text read
        return parse_txt(file_path)
